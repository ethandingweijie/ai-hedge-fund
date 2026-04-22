"""
update_reference_docx.py — appends v1.8 + v1.9 changelog sections to
AI_Hedge_Fund_Reference.docx, matching the existing section style, and
bumps the version header.

Run:
    python -m scripts.update_reference_docx
"""
from __future__ import annotations
import os, sys
_REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if _REPO_ROOT not in sys.path:
    sys.path.insert(0, _REPO_ROOT)

from docx import Document

DOCX_PATH = r"C:\Users\ethan\Documents\Projects\AI Hedge Fund\AI_Hedge_Fund_Reference.docx"

def bump_version(doc):
    """Bump the version header at the top from 'Version 1.7 | 20 April 2026' to v1.9."""
    for p in doc.paragraphs[:10]:
        if "Version 1.7" in p.text or "Version 1.8" in p.text:
            for run in p.runs:
                run.text = run.text.replace("Version 1.7", "Version 1.9")
                run.text = run.text.replace("Version 1.8", "Version 1.9")
                run.text = run.text.replace("20 April 2026", "22 April 2026")
                run.text = run.text.replace("21 April 2026", "22 April 2026")
            print(f"[ OK ] bumped header: '{p.text}'")
            return


def add_paragraph(doc, text: str, style: str = "Normal") -> None:
    """Append a paragraph with the given style."""
    p = doc.add_paragraph(text, style=style)
    return p


def main():
    doc = Document(DOCX_PATH)

    # ── Bump version header ──
    bump_version(doc)

    # ── Append v1.8 changelog ──
    doc.add_paragraph(
        "2026-04-22 — v1.8: Tier 2 Sector Methodology Rebuild "
        "(Banks, REITs, Biopharma, Tech/SaaS)",
        style="Heading 2",
    )

    doc.add_paragraph("Tier 2 Bank Methodology (institutional rebuild)", style="Heading 3")
    for b in [
        "2-stage Residual Income replaces primitive ROE-CoE spread. ROE fades linearly current → profile "
        "target over 5-10 years; BVPS compounds at retention × ROE; terminal spread (+50-100 bps moat "
        "premium for GSIBs / Super-Regionals / Indian privates) captures durable excess returns.",
        "P/TBV replaces P/BV — strips goodwill + intangibles. Now correctly uses FMP's direct "
        "tangibleBookValuePerShare field from /stable/ratios (fixes earlier double-count of goodwill "
        "that had JPM TBV/sh reported as $90.81 instead of $106.85).",
        "CET1 Excess Capital overlay — CET1 > target returns (actual-target) × RWA × 0.70 per share "
        "(asymmetric haircut: only 70% distributable); CET1 < target subtracts deficit at full haircut. "
        "RWA proxied via sub-profile-specific asset ratios when FMP lacks regulatory disclosures.",
        "P/E (norm) through-cycle fallback — uses equity × target_ROE instead of trailing NI. Immune "
        "to credit-cycle provision distortion.",
        "Buyback-aware retention rate — includes common_stock_repurchased alongside dividends (JPM "
        "returned ~$25B via repurchases in 2024; dividend-only retention overstated by 30+ pp).",
        "Profile weights flipped: RI 55% / P/TBV 25% / P/E (norm) 15% / Excess Capital 5%. Dropped "
        "'ROE vs CoE' (double-counted RI per Gemini critique).",
        "10 bank sub-profiles with geography-aware calibration: Money Center (US/EU), Regional, "
        "Super-Regional, EM Bank (China SOE), EM Bank Premium (India private, 7y fade + 16% target "
        "ROE), Investment Bank, Mortgage/GSE, Neo/Challenger (10y J-curve fade), Brokerage.",
        "Deep research _extract_bank_metrics() LLM extractor — CET1, NIM, efficiency ratio, NPL, "
        "NPL coverage, management target ROE/ROTCE, loan/deposit growth, management overlays, "
        "NIM rate sensitivity, forward guidance quotes.",
    ]:
        doc.add_paragraph(b, style="List Paragraph")

    doc.add_paragraph("Tier 2 REIT NAV / P/FFO / P/AFFO (replaces all-proxy)", style="Heading 3")
    for b in [
        "NAV (Cap Rates): NOI / cap_rate − total_debt + cash. Scenario-invariant anchor.",
        "P/FFO + P/AFFO — sub-type-specific multiples, replacing prior P/E proxy (GAAP earnings "
        "depressed by non-cash real-estate D&A).",
        "AFFO-gated DDM — clamps dividends to AFFO/share, catches yield-trap valuations.",
        "12 REIT sub-types with maintenance capex caps: data_center 2% / lab 2.5% / industrial 3% / "
        "self_storage 3% / net_lease 1% / residential 4% / healthcare 4% / retail 5.5% / office 6% / "
        "hospitality 7.5% / infrastructure 8.5% of revenue. net_lease added 2026-04-22 after Gemini "
        "critique on Realty Income (O NAV $24.55 → $42.66).",
        "Deep research _extract_reit_metrics() — cited cap rate, occupancy, WALE, sub-type/geo mix, "
        "DPU vs AFFO coverage, leverage. Overrides sub-type defaults via cap_rate_market.",
    ]:
        doc.add_paragraph(b, style="List Paragraph")

    doc.add_paragraph("Tier 2 Biopharma rNPV", style="Heading 3")
    for b in [
        "2-stage rNPV: per-asset peak_sales × op_margin × (1-tax) × ramp_profile × cumulative_PoS × "
        "discount(years_to_launch).",
        "PHASE_POS_TABLE: Ph1 9.6%, Ph2 15.3%, Ph3 49.3%, Filed 85%, Approved 100% (BIO 2011-2020 "
        "industry stats + FDA historical).",
        "Therapeutic-area PoS multipliers — Oncology 0.55x, CNS 0.60x, Rare 1.7x, Hematology 1.4x, "
        "GLP-1 1.30x.",
        "Bell-shaped commercial stream (20/50/80% ramp + 7 years peak + 40/20/10% LOE decay).",
        "_extract_pipeline_assets() extractor from deep research sections 2A/2D/2F — per-asset JSON "
        "with name, phase, peak_sales, launch_year, indication, evidence.",
    ]:
        doc.add_paragraph(b, style="List Paragraph")

    doc.add_paragraph("Tier 3 Insider-Activity WACC Overlay", style="Heading 3")
    for b in [
        "Net 12m insider buying / selling translates to ±bp WACC modifier (capped at ±50 bp). "
        "Cluster buys get additional tightening; CEO/CFO conviction sells widen.",
    ]:
        doc.add_paragraph(b, style="List Paragraph")

    # ── Append v1.9 changelog — frontend panels + deploy fixes ──
    doc.add_paragraph(
        "2026-04-22 — v1.9: Sector-Specific Valuation UI + TBV / NAV Calibration Fixes",
        style="Heading 2",
    )

    doc.add_paragraph("REIT Valuation Panel (frontend)", style="Heading 3")
    for b in [
        "Ships app/frontend/src/components/report/reit/REITValuationPanel.tsx — 8 sector-specific "
        "sub-panels mirroring institutional REIT research conventions.",
        "NAV Hero card: centered NAV/sh ($167.79 for DLR · $42.66 for O) with upside vs current "
        "price; quad grid below (NOI, GAV, Debt tinted red, Cash tinted green).",
        "REIT Key Stats grid: Implied cap, Dist. yield, AFFO coverage, Leverage, Occupancy, WALE, "
        "FFO/sh, AFFO/sh with threshold color-coding.",
        "Distribution Quality gauge + 100% safety line for AFFO coverage.",
        "NPI + DPU history bar charts (5y, CLINT-style) via recharts.",
        "Portfolio Composition pies (by asset class + geography) — always rendered, falls back to "
        "classified sub-type as 100% slice when deep-research extractor hasn't populated the mix.",
        "Cap-Rate × NOI-Growth 3×3 sensitivity matrix, peer cell highlighted.",
        "Backend emitter dcf_range.reit_breakdown with full data structure for 5y history arrays.",
    ]:
        doc.add_paragraph(b, style="List Paragraph")

    doc.add_paragraph("Bank Valuation Panel (frontend)", style="Heading 3")
    for b in [
        "Ships app/frontend/src/components/report/bank/BankValuationPanel.tsx — 8 panels matching "
        "DBS / OCBC / Gemini institutional research driver hierarchy.",
        "P/TBV Fair Value Hero (Gordon-growth identity): Fair = TBV × (1 + (ROE-CoE) / CoE). "
        "JPM: $186 at TBV $106.85 × 1.74x; GS: $472 at TBV $377.94 × 1.25x.",
        "Quad grid: TBV/sh, BVPS, ROE (tinted green ≥ target), CET1 buffer bps (tinted).",
        "Bank Key Stats 8-tile grid: ROE, ROA, NIM, CIR, Credit Cost (falls back to NCO), BVPS, "
        "NPL, CET1 — with profile-calibrated threshold color-coding.",
        "ROE vs CoE Spread gauge — horizontal bar, zero-line marker, green fill for positive spread.",
        "Capital Return card — total yield hero (div + buyback); 4-tile row with div yield, "
        "buyback yield, payout ratio, CET1 surplus tinted green when positive.",
        "Pre-Provision Operating Profit 5y bar chart — via _compute_ppop() with 3-tier fallback.",
        "NIM History 5y + appended NIM Rate Sensitivity tile (shows 'X bps NIM per 100 bps rate').",
        "Loan Growth card — 5y bars when FMP exposes loan book; else single YoY tile from research.",
        "Book Quality card — NPL, NPL coverage ratio (with 100% safety gauge), credit cost, "
        "management overlays. Only renders when ≥1 research-extracted field is present.",
    ]:
        doc.add_paragraph(b, style="List Paragraph")

    doc.add_paragraph("Sector routing & data integrity", style="Heading 3")
    for b in [
        "21 major US REITs + 6 net-lease REITs added to TICKER_SECTOR_LOOKUP for deterministic "
        "routing: DLR, EQIX, PSA, EXR, ARE, WELL, VTR, AVB, EQR, MAA, ESS, VICI, BXP, VNO, STAG, "
        "HST, RHP, APLE, KIM, FRT, REG, MAC, DOC, OHI, ADC, NNN, WPC, SRC, BNL.",
        "SGX bank profiles upgraded to Money Center Bank: O39.SI (OCBC), U11.SI (UOB) joined "
        "D05.SI (DBS) with proper sub-profile calibration.",
        "FMP _BALANCE_MAP corrected — intangibleAssets now maps to intangible_assets (was mapping "
        "goodwillAndIntangibleAssets which double-counted goodwill when stripped). Cascades to "
        "every sector but matters most for banks where TBV is a primary valuation anchor.",
        "FMP _RATIOS_MAP extended with tangibleBookValuePerShare as preferred TBV/sh source "
        "(applies bank's own reporting convention — e.g. JPM treats MSRs as tangible).",
        "FMP _BALANCE_MAP extended with netLoans, loansAndLeasesReceivables, loansHeldForInvestment, "
        "totalDeposits for future bank loan-book coverage.",
    ]:
        doc.add_paragraph(b, style="List Paragraph")

    doc.add_paragraph("12-Month Price Target methodology — REITs", style="Heading 3")
    for b in [
        "REITs were previously routed through the _use_pe_only branch applying EPS × 35x (RealEstate "
        "peer PE). This is conceptually wrong because REIT GAAP EPS is heavily depressed by non-cash "
        "real-estate D&A — EPS/FFO ratio varies 0.3-0.5 across sub-types.",
        "New REIT branch: FFO/sh × (1+g) × P/FFO_sub-type blended 60/40 with AFFO/sh × (1+g) × "
        "P/AFFO_sub-type. No growth_premium on top — sub-type multiples already embed growth. "
        "Calibration: DLR base case PT moves from $136 to $237, matching 22x P/FFO_fwd market multiple.",
        "Scenario multipliers (0.75 / 1.00 / 1.25) remain the sole dispersion mechanism across "
        "bear / base / bull.",
    ]:
        doc.add_paragraph(b, style="List Paragraph")

    doc.add_paragraph("Gemini review remediation", style="Heading 3")
    for b in [
        "JPM TBV / Fair Value — Gemini flagged our $90.81 TBV/sh vs JPM's reported Q4'25 $107.56. "
        "Root cause was goodwill double-count in _BALANCE_MAP mapping; fix uses FMP's direct "
        "tangibleBookValuePerShare ratio. JPM Fair Value recomputed $158.08 → $186.00 (within 1% "
        "of Gemini's $187.15 reconciliation).",
        "Realty Income (O) NAV — Gemini flagged our $24.55 NAV/sh vs O's $40-45 Book Value range. "
        "Root cause was 'default' sub-type (6.5% cap rate); new net_lease sub-type (5.0% cap rate, "
        "16x P/FFO, 18x P/AFFO) corrects O NAV/sh $24.55 → $42.66. Cascades to ADC, NNN, WPC, BNL.",
        "12M PT methodology — Gemini flagged EPS × P/E branch for REITs. Replaced with P/FFO + "
        "P/AFFO blend (see above). Gemini's static Gordon growth critique on banks (missing g term) "
        "deferred to v1.10 — current RI 2-stage model already handles compounding.",
    ]:
        doc.add_paragraph(b, style="List Paragraph")

    doc.add_paragraph("Backfill tooling", style="Heading 3")
    for b in [
        "scripts/backfill_reit_breakdown.py — re-derives reit_breakdown from line items for "
        "archived runs predating the REIT panel. Targets web_runs.full_result_json (the table "
        "the deployed app reads).",
        "scripts/backfill_bank_breakdown.py — same pattern for banks.",
        "POST /admin/backfill-reit-breakdown and /admin/backfill-bank-breakdown — one-shot HTTP "
        "endpoints gated behind DB_UPLOAD_SECRET, callable via curl / Invoke-RestMethod without "
        "shell access to Railway. Dry-run default; supports ticker filter and force re-derive.",
        "Ticker-whitelist fallback: when sector column is null or mis-classified on archived rows, "
        "backfill cross-references TICKER_SECTOR_LOOKUP and auto-corrects sector='RealEstate' / "
        "'Financials' on patch.",
    ]:
        doc.add_paragraph(b, style="List Paragraph")

    doc.add_paragraph("Frontend wiring", style="Heading 3")
    for b in [
        "REIT + Bank panels wired into all 3 render paths: pages/ReportPage.tsx (live runs), "
        "pages/ReportViewPage.tsx (historic desktop), components/v2/V2ReportView.tsx (mobile).",
        "Gate ordering: REIT > Bank > generic DCF Ladder. Non-applicable sectors fall through "
        "to the existing ladder untouched.",
        "Zinc palette alignment: swapped shadcn bg-card / border-border / text-foreground tokens "
        "to explicit bg-white dark:bg-zinc-900 / border-zinc-200 dark:border-zinc-800 / "
        "text-zinc-900 dark:text-zinc-50 to match v2 card surfaces exactly. Fixes grey "
        "intermediate-shade bug that made REIT / Bank panels visually distinct from the rest.",
        "TypeScript types: lib/reportTypes.ts extended with ReitBreakdown and BankBreakdown "
        "interfaces on DcfRange.",
    ]:
        doc.add_paragraph(b, style="List Paragraph")

    doc.add_paragraph("Build numbers", style="Heading 3")
    for b in [
        "Bumped app/backend/main.py FastAPI version to 1.9.0. Bumped app/frontend/package.json "
        "to 1.9.0. Bumped pyproject.toml to 1.9.0.",
    ]:
        doc.add_paragraph(b, style="List Paragraph")

    # Save
    doc.save(DOCX_PATH)
    print(f"[ OK ] appended v1.8 + v1.9 sections to {DOCX_PATH}")
    print(f"       final paragraph count: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
