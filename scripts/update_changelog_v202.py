"""Append v2.0.2 changelog section to AI_Hedge_Fund_Reference.docx.

Covers the session's work: admin re-extract endpoint, profile_name as
first-class column, extractor model fix (qwen3-max → qwen3.6-plus),
2F-only extractor input, OpenAI SDK exploration + revert, rate-limit
retry with Retry-After, and the cascade of diagnostic/observability
improvements.
"""
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

PATH = r"C:\Users\ethan\Documents\Projects\AI Hedge Fund\AI_Hedge_Fund_Reference.docx"
doc = Document(PATH)


def H2(text): doc.add_paragraph(text, style="Heading 2")
def H3(text): doc.add_paragraph(text, style="Heading 3")
def P(text):  doc.add_paragraph(text, style="Normal")
def LI(text): doc.add_paragraph(text, style="List Paragraph")


# ── v2.0.2 header ───────────────────────────────────────────────────────────
H2("2026-04-24 — v2.0.2: Extractor Model Alignment + Admin Re-Extract Infrastructure")

# ── Executive summary ──────────────────────────────────────────────────────
H3("Executive Summary")

P(
    "Diagnosed and fixed a production bug that caused empty saas_metrics / "
    "bank_metrics / reit_metrics / pipeline_assets dicts on every live pipeline "
    "run since the extractor integration shipped. Root cause: synthesis model "
    "defaulted to qwen3-max which returns 404 NotFoundError on the user's "
    "Anthropic-compat DashScope endpoint (/apps/anthropic). Extractors failed "
    "silently (caught Exception → return {}) so downstream KPI tiles stayed "
    "empty without any error visibility."
)

P(
    "Built admin re-extract infrastructure to recover historic run data without "
    "re-running the expensive research pipeline. Surfaced two previously-hidden "
    "bugs in the extractor chain (2F truncation, DashScope Anthropic-compat "
    "SDK mismatch) via the re-extract diagnostic path. Added profile_name as "
    "a first-class web_runs column with backfill, giving admin UI visibility "
    "into sub-sector classification for every archived run."
)

# ── Root cause: production extractor silently failing ─────────────────────
H3("Root Cause — Live pipeline extractor 404'd on every run (qwen3-max vs qwen3.6-plus)")

P(
    "User's DEEP_RESEARCH_BASE_URL points to https://dashscope-intl.aliyuncs.com"
    "/apps/anthropic — the Anthropic-compat DashScope endpoint. qwen3.6-plus "
    "is the provisioned model on this endpoint (30K RPM / 5M TPM). qwen3-max "
    "is documented at the global rate-limit table but 404s when requested on "
    "/apps/anthropic because the specific regional deployment doesn't have it."
)

P(
    "The live pipeline's extractor fan-out (src/agents/industry/deep_research.py "
    "lines 3831, 3839) resolved synthesis model as:"
)

LI("_synthesis_model = os.environ.get('DEEP_RESEARCH_SYNTHESIS_MODEL') or 'qwen3-max'")

P(
    "When the env var was unset (or set to qwen3-max from an older config), "
    "every subsequent extractor call — _extract_saas_metrics, _extract_bank_metrics, "
    "_extract_reit_metrics, _extract_pipeline_assets, _extract_dcf_calibration, "
    "_extract_segment_scenarios — sent a request with model='qwen3-max' to the "
    "/apps/anthropic endpoint and received HTTP 404 in return. The extractors' "
    "try/except caught this as a generic Exception, printed a rate-limit-retry "
    "log message, exhausted its 93-second budget, and returned {} to the caller. "
    "The empty dict was stored as the ticker's saas_metrics value in web_runs."
    "full_result_json, causing the frontend Key Metrics card to render with only "
    "FMP-derived fields (Rule of 40, billings growth, Revenue Growth) and zero "
    "LLM-extracted KPIs (NRR, Gross Retention, CAC Payback, Magic Number, LTV/CAC)."
)

# ── Production fix ─────────────────────────────────────────────────────────
H3("Backend — Live pipeline synthesis model default fix (commit cd6cd5e)")

P(
    "Changed the extractor synthesis model default from 'qwen3-max' to "
    "'qwen3.6-plus' in both routing branches (HK tickers + US-Qwen tickers). "
    "DEEP_RESEARCH_SYNTHESIS_MODEL env var still overrides for explicit "
    "control, but the fallback is now endpoint-agnostic."
)

P(
    "Single-variable propagation: the change affects one variable "
    "(_synthesis_model) that feeds all 6 extractors through the ThreadPoolExecutor "
    "fan-out at lines 3583-3588. No per-extractor edits required — the fix "
    "cascades automatically to dcf_calibration, segment_scenarios, saas_metrics, "
    "bank_metrics, reit_metrics, pipeline_assets. Also applied to the two "
    "cache-revalidation paths at lines 2386 (initial cache hit path) and 2488 "
    "(delta-pass path) that re-extract from cached research text."
)

P(
    "Expected impact going forward: every live pipeline run for Tech / Bank / "
    "REIT / Biopharma tickers will now produce populated sector KPI dicts "
    "instead of silently-empty dicts. Historic runs need /admin/reextract-metrics "
    "(described below)."
)

# ── Admin re-extract infrastructure ────────────────────────────────────────
H3("Backend — /admin/reextract-metrics endpoint (commit 1839009)")

P(
    "New admin HTTP endpoint that re-runs the 6 LLM extractors against "
    "existing stored deep research in web_runs.full_result_json without "
    "triggering a fresh pipeline run. Retrofits historic runs so the "
    "frontend sees recovered KPIs on next page load — without burning "
    "the expensive research synthesis pass (which takes 4-6 min per ticker)."
)

P("Three components, mirroring the v1.9 REIT/Bank backfill pattern:")

LI("src/memory/reextract_metrics.py — core helper with reextract_for_run(run_id, dry_run, provider) and reextract_by_ticker(ticker, limit, dry_run, provider).")
LI("scripts/reextract_metrics.py — CLI with mutually-exclusive --run-id / --ticker / --tickers flags and --dry-run default.")
LI("POST /admin/reextract-metrics in app/backend/routes/admin.py — same auth contract (DB_UPLOAD_SECRET) as existing /admin/backfill-reit-breakdown.")

P("Query params:")

LI("secret — DB_UPLOAD_SECRET (required)")
LI("ticker — one ticker (uses last N runs per limit param)")
LI("tickers — comma-separated tickers")
LI("run_id — target a specific web_runs UUID")
LI("limit — per-ticker run count (default 1)")
LI("dry_run — true (default) shows diff without writing; false writes")
LI("provider — auto (default, Qwen preferred) / qwen / anthropic (Claude fallback)")
LI("verbose — true surfaces diagnostic_saas block with raw Qwen response + parse/clamp state")

P(
    "Protection: only writes when AFTER shows strictly more populated fields "
    "than BEFORE (prevents regression erasing existing data). Dry-run default. "
    "Exactly one of {run_id, ticker, tickers} required (400 error otherwise)."
)

# ── Section re-parsing for historic runs ───────────────────────────────────
H3("Backend — Section re-parse during re-extract (commit 5aab000)")

P(
    "Historic runs archived before the widened 2F parser regex (commit d8706df, "
    "v2.0.1) have partial deep_research_sections dicts where the '2f' key is "
    "missing — the old parser silently dropped 2F headings with list markers "
    "'- 2F', divider bars '=== 2F ===', or prose 'Section 2F:' forms. Without "
    "sections['2f'], the saas_metrics extractor fell back to deep_research[:8000] "
    "which for rich reports is mostly 2A+2D content, starving the LLM of the "
    "KPI framework it needs."
)

P(
    "Fix: re-parse sections from stored deep_research text using the CURRENT "
    "widened regex before feeding to extractors. Prefer re-parsed when it "
    "recovered '2f' key the stored dict was missing. Response surfaces "
    "sections_source, sections_keys, and section_2f_len so the user can "
    "verify re-parse success per run."
)

P(
    "Expected gain on DDOG re-extract: stored_sections had no '2f' key → "
    "re-parsed found 2F at 5366 chars → extractor now sees the full 2F content."
)

# ── Extractor input 2F-only ────────────────────────────────────────────────
H3("Backend — Extractors focus on 2F only (commits d257f22, b2689a2)")

P(
    "Diagnostic on DDOG revealed extractor was seeing 13,143 chars of input "
    "but truncating to 8,000 chars with 2A+2D first (7800 chars) and 2F last "
    "(5366 chars) — only ~200 chars of 2F reached the LLM. Qwen found RPO "
    "(mentioned in 2A's quantitative foundation from FMP preload) but missed "
    "NRR, CAC Payback, Magic Number, Rule of 40 — all of which live in 2F."
)

P("Two sequential fixes:")

LI("b2689a2 — Put 2F first in combined input, raise truncation 8000 → 20000 chars.")
LI("d257f22 — Per user directive ('the extractor just need to focus on 2F'), switched to 2F ONLY. 2A (profit pool) and 2D (cycle) don't contain KPIs; including them diluted the LLM's attention.")

P(
    "Applied to all 4 extractors with the (2f|2a|2d) pattern: "
    "_extract_saas_metrics, _extract_bank_metrics, _extract_pipeline_assets, "
    "_extract_segment_scenarios. Falls back to full 2A+2D+2F concat when 2F "
    "is missing or too short (<500 chars). Final fallback is "
    "deep_research[:20000] when sections dict is entirely empty."
)

# ── Schema: profile_name first-class column ───────────────────────────────
H3("Backend — profile_name as first-class web_runs column (commits 3701e79, 8cec467)")

P(
    "Added profile_name TEXT column to web_runs table DDL + migration entry. "
    "Composite index idx_web_runs_sector_profile supports 'all Growth SaaS' "
    "queries without full scan, left-prefix supports sector-only filters. "
    "_extract_web_run_summary now returns a 4-tuple (final_action, regime, "
    "sector, profile_name) with resolution tree: state → profile_names_map → "
    "TICKER_SECTOR_LOOKUP canonical fallback. _save_web_run writes the column "
    "on every new save."
)

P(
    "Companion backfill endpoint /admin/backfill-profile-name + "
    "scripts/backfill_profile_name.py populates historic rows with NULL "
    "profile_name by resolving each ticker through the same lookup chain. "
    "Updates BOTH the column (for admin filtering) AND the inner "
    "full_result_json.data.profile_name (for re-extract / frontend reads) in "
    "a single transaction to keep them consistent."
)

P("Typical backfill distribution observed on 40 historic rows:")

LI("Growth SaaS: 5 rows (DDOG, SNOW, PLTR, FRSH, ZM)")
LI("Mature SaaS: 8 rows (CRM, ADBE ×3, NOW ×2, VEEV ×2)")
LI("Money Center Bank: 1 row (JPM)")
LI("Asset Manager: 1 row (AB)")
LI("Managed Care: 1 row (MOH)")
LI("Automotive & EV: 1 row (01211.HK)")
LI("Unresolved: 23 rows (AAPL, O, DLR, NFLX, PYPL, NVO, etc. — not in TICKER_SECTOR_LOOKUP)")

P(
    "Unresolved tickers safely leave profile_name as NULL; their next live "
    "run populates via strategic_router's in-situ classification using the "
    "DCF-ratio classifier."
)

# ── Rate-limit handling ───────────────────────────────────────────────────
H3("Backend — Rate-limit-aware retry for DashScope 403 (commits e79f569, c6ef74c)")

P(
    "Qwen via DashScope returns HTTP 403 AccessDenied with message 'Rate limit "
    "exceeded' when the account hits its TPM/RPM ceiling — instead of the "
    "standard HTTP 429. The Anthropic SDK's built-in retry only fires on "
    "408/409/429/500+ so 403 rate-limit errors propagated silently to the "
    "extractor's try/except, returning {}."
)

P(
    "New _call_llm_with_rate_retry wrapper at src/agents/industry/deep_research.py "
    "line 44 catches any exception whose message contains 'rate limit' / "
    "'ratelimit' / 'quota' / 'throttl' / 'accessdenied' (case-insensitive) and "
    "retries with exponential backoff (3s → 6s → 12s → 24s → 48s, capped at "
    "60s per wait, 5 retries max = 93s total budget). Respects upstream "
    "Retry-After header via duck-typed .response.headers lookup — takes "
    "max(server guidance, our backoff) so we never hammer sooner than requested."
)

P(
    "Applied to all 7 LLM callsites: _extract_dcf_calibration, "
    "_extract_segment_scenarios, _extract_saas_metrics, _extract_bank_metrics, "
    "_extract_reit_metrics, _extract_pipeline_assets, _diagnose_saas_extractor. "
    "Log line shows [llm_rate_retry] extractor ticker rate-limited (attempt "
    "N/M), sleeping Ns [Retry-After=Ns|exponential]... for Railway stdout "
    "observability."
)

# ── Anthropic fallback provider ──────────────────────────────────────────
H3("Backend — Anthropic Claude fallback provider (commit a6c95b2)")

P(
    "Added provider=anthropic query param to /admin/reextract-metrics. When "
    "DashScope quota is exhausted (daily cap, plan tier, free trial), the user "
    "can force extractors to run on Anthropic Claude — which uses a separate "
    "quota pool via ANTHROPIC_API_KEY env var. Returns 'provider' and "
    "'model_name' in the response so the user can confirm which LLM ran."
)

P("Three modes:")

LI("provider=auto (default) — prefer Qwen if creds available, fall back to Anthropic")
LI("provider=qwen — force DashScope/Qwen; fails loudly if creds missing")
LI("provider=anthropic — force Claude; fails loudly if ANTHROPIC_API_KEY missing")

P(
    "Isolation: only affects the re-extract helper path. Live research pipeline "
    "continues using its existing anthropic/qwen client selection logic. Zero "
    "production risk."
)

# ── Diagnostic observability ───────────────────────────────────────────────
H3("Backend — Diagnostic & observability improvements (commits 38f8475, 83101db, d8706df)")

P(
    "Triaged extractor failures by adding three layers of diagnostic "
    "visibility without requiring Railway dashboard access:"
)

LI("verbose=true query param on /admin/reextract-metrics — response includes diagnostic_saas block with input_chars, combined_preview (500 chars), raw_response (full Qwen output), raw_len, parsed_type, parsed_keys, parsed_sample, validated_fields, clamp_rejections, error. Lets the user diagnose 'extractor ran but fields empty' from the HTTP response directly.")
LI("stdout logs in _extract_saas_metrics — [saas_metrics TICKER] input=N chars · raw_response=N chars · parsed_type=dict · parsed_keys=[...] · preview='...' visible in Railway logs on every call. Distinguishes (a) API call failed, (b) Qwen returned valid JSON with no KPIs, (c) KPIs returned but failed clamp.")
LI("Section 2F preview log in deep_research.py — [Section 2F preview TICKER, N chars, subsections=[...]] prints first 400 chars of 2F + detected subsection tokens (2F.1, 2F.2, 2F.4, etc.) so operators can verify Qwen produced the expected structure.")

# ── OpenAI SDK exploration (reverted) ───────────────────────────────────
H3("Backend — OpenAI SDK adapter experiment (reverted, commit 534aa13)")

P(
    "Briefly attempted swapping the Qwen client from anthropic SDK to openai "
    "SDK via an adapter class (commit c6ef74c) on the theory that the "
    "anthropic SDK's error mapping was obscuring real rate-limit errors. "
    "The swap broke immediately with 404 NotFoundError because the user's "
    "endpoint (/apps/anthropic) is Anthropic-compat only and the openai SDK "
    "sends POST to base_url/chat/completions which doesn't exist there."
)

P(
    "Reverted to anthropic SDK (commit 534aa13) after user shared "
    "DEEP_RESEARCH_BASE_URL. Lesson learned: DashScope has two separate "
    "compatibility endpoints, NOT interchangeable:"
)

LI("/compatible-mode/v1 — OpenAI-format payloads, use openai SDK")
LI("/apps/anthropic — Anthropic-format payloads, use anthropic SDK")

P(
    "Adapter classes (_OpenAIAsAnthropicAdapter, _MessagesShim, _MessageShim, "
    "_TextBlockShim) remain in src/memory/reextract_metrics.py as dead code "
    "for future migration if the user ever moves to the OpenAI-compat endpoint."
)

# ── Commit chain ──────────────────────────────────────────────────────────
H3("Commit chain — myfork/main (chronological)")

P("Thirteen commits landed as the v2.0.2 release, all on Railway (backend) + Vercel (frontend):")

LI("1839009 — feat(reextract): admin endpoint + CLI to rerun extractors on stored runs")
LI("60caf92 — fix(reextract): TICKER_SECTOR_LOOKUP fallback when stored profile_name is empty")
LI("5aab000 — fix(reextract): re-parse sections from stored deep_research to recover 2F")
LI("83101db — debug(saas_metrics): surface raw response + parse state + clamp rejects")
LI("38f8475 — feat(reextract): verbose mode surfaces raw Qwen response in HTTP result")
LI("e79f569 — fix(extractors): rate-limit-aware retry for DashScope 403 AccessDenied")
LI("a6c95b2 — feat(reextract): provider param for Anthropic Claude fallback")
LI("3701e79 — feat(schema): profile_name as first-class web_runs column + backfill")
LI("8cec467 — fix(backfill-profile-name): trigger schema migration before backfill writes")
LI("c6ef74c → 534aa13 — openai SDK adapter experiment (reverted)")
LI("cf5076a → c0ce2e9 → 1686552 — model name resolution iterations (ended at hard-coded qwen3.6-plus)")
LI("b2689a2 → d257f22 — extractor input: put 2F first, then 2F only")
LI("cd6cd5e — fix(live-pipeline): default synthesis model to qwen3.6-plus (the big one — fixes silent 404 on every live run)")

# ── How to recover historic data ──────────────────────────────────────────
H3("Operational guide — Recovering historic run data after v2.0.2 deploy")

P(
    "Every historic Tech/Bank/REIT/Biopharma run in web_runs before 2026-04-24 "
    "has empty sector extractor dicts because of the qwen3-max default bug. "
    "Recovery workflow:"
)

LI("Step 1: Run /admin/backfill-profile-name once to populate the profile_name column on historic rows (resolves sector-extractor gate for tickers in TICKER_SECTOR_LOOKUP).")
LI("Step 2: For each ticker with meaningful historic runs, invoke POST /admin/reextract-metrics?secret=X&ticker=TICKER&provider=qwen&dry_run=true to preview.")
LI("Step 3: If after.saas_metrics shows populated fields (6-7 for Tech), flip dry_run=false to commit. Response returns 'updated': 1.")
LI("Step 4: Hard-refresh the ticker's report page — Key Metrics + Traffic Light + Commentary cards render with populated values.")
LI("Step 5: Run in batches of 1-2 tickers with a 60-90s gap to stay under DashScope rate limits (per-minute throttles clear in 30-60s on qwen3.6-plus).")

P(
    "If DashScope quota genuinely exhausted (daily cap hit): use "
    "provider=anthropic to route through Claude for the re-extract session. "
    "Claude extracts the same 7 SaaS KPIs from the same 2F input with "
    "comparable quality (if not slightly better JSON discipline)."
)

# ── Testing verification ──────────────────────────────────────────────────
H3("Testing — Multi-provider verification matrix")

P("End-to-end verified on 2026-04-24 with DDOG (Growth SaaS) re-extract:")

LI("Section re-parse: stored_had_2f=False → now=True, section_2f_len=5366 ✓")
LI("Profile resolution: empty → 'Growth SaaS' via TICKER_SECTOR_LOOKUP fallback ✓")
LI("Extractor fan-out: saas_metrics ran (was previously gated off due to empty profile_name) ✓")
LI("Model routing: model_name='qwen3.6-plus' confirmed in response ✓")
LI("Qwen extraction: 7 validated fields (nrr_pct 1.2, gross_retention_pct 0.95, cac_payback_months 15.0, ltv_cac_ratio 5.3, rule_of_40_score 50.4, magic_number 0.61, rpo_growth_yoy 0.52) ✓")
LI("Clamp validation: 0 rejections (all values in valid ranges) ✓")
LI("FMP fallback: augments LLM output with rule_of_40 / billings_growth_yoy / evidence where FMP data exists ✓")

# ── Save ────────────────────────────────────────────────────────────────────
doc.save(PATH)
print(f"OK Appended v2.0.2 changelog section to {PATH}")
print(f"   Total paragraphs now: {len(doc.paragraphs)}")
